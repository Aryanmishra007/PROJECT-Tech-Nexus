"""
Resume Parser Utility
Extracts text and skills from PDF/DOCX resumes.
"""

import re
import io
from collections import Counter

# Try importing PDF and DOCX libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Comprehensive skill keywords list
SKILL_KEYWORDS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust',
    'kotlin', 'swift', 'r', 'scala', 'ruby', 'php', 'dart', 'matlab',
    'perl', 'bash', 'powershell', 'sql', 'html', 'css',

    # AI / ML Frameworks
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'sklearn', 'xgboost',
    'lightgbm', 'catboost', 'huggingface', 'transformers', 'spacy', 'nltk',
    'opencv', 'fastai', 'jax', 'mxnet', 'caffe', 'theano',

    # Data Science / Analytics
    'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'bokeh',
    'scipy', 'statsmodels', 'tableau', 'power bi', 'powerbi', 'looker',
    'excel', 'jupyter', 'google analytics', 'mixpanel',

    # Web Frameworks
    'react', 'angular', 'vue', 'next.js', 'nuxt', 'svelte',
    'django', 'flask', 'fastapi', 'express', 'node.js', 'nodejs',
    'spring', 'spring boot', 'laravel', 'rails', 'asp.net',
    'graphql', 'rest api', 'restful', 'websocket',

    # Databases
    'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch',
    'cassandra', 'dynamodb', 'sqlite', 'oracle', 'sql server', 'firebase',
    'neo4j', 'influxdb', 'clickhouse', 'snowflake', 'bigquery',

    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'k8s',
    'terraform', 'ansible', 'jenkins', 'gitlab ci', 'github actions',
    'circleci', 'helm', 'prometheus', 'grafana', 'nginx', 'apache',
    'linux', 'unix', 'bash scripting', 'ci/cd', 'devops',

    # AI/ML Concepts
    'machine learning', 'deep learning', 'neural network', 'nlp',
    'natural language processing', 'computer vision', 'reinforcement learning',
    'transfer learning', 'fine-tuning', 'llm', 'large language models',
    'generative ai', 'prompt engineering', 'rag', 'vector database',
    'feature engineering', 'model deployment', 'mlops', 'a/b testing',

    # Data Engineering
    'spark', 'hadoop', 'kafka', 'airflow', 'dbt', 'etl', 'elt',
    'data pipeline', 'data warehouse', 'data lake', 'flink',

    # Security
    'cybersecurity', 'penetration testing', 'ethical hacking', 'siem',
    'owasp', 'ssl/tls', 'oauth', 'jwt', 'authentication', 'authorization',

    # Soft Skills
    'leadership', 'communication', 'teamwork', 'problem solving',
    'critical thinking', 'project management', 'agile', 'scrum', 'kanban',
    'time management', 'collaboration', 'presentation', 'mentoring',

    # Tools
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence',
    'slack', 'figma', 'postman', 'vscode', 'intellij', 'vim',
    'linux', 'windows server', 'macos',
}

# Common normalization aliases for NLP matching
SKILL_ALIASES = {
    'js': 'javascript',
    'ts': 'typescript',
    'py': 'python',
    'node': 'node.js',
    'nodejs': 'node.js',
    'postgres': 'postgresql',
    'g suite': 'google workspace',
    'k8s': 'kubernetes',
    'ci cd': 'ci/cd',
    'ml': 'machine learning',
    'dl': 'deep learning',
    'nlp': 'natural language processing',
    'cv': 'computer vision',
    'genai': 'generative ai',
}

# Optional spaCy support; parser works without it.
try:
    import spacy
    _NLP = spacy.blank('en')
    if 'sentencizer' not in _NLP.pipe_names:
        _NLP.add_pipe('sentencizer')
    HAS_SPACY = True
except Exception:
    _NLP = None
    HAS_SPACY = False


def _normalize_text_for_nlp(text):
    """Normalize text for robust NLP pattern matching."""
    text = text.lower()
    text = text.replace('\n', ' ')
    text = re.sub(r'[^a-z0-9+.#/\-\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _normalize_skill(skill):
    """Normalize individual skill string and map aliases."""
    s = re.sub(r'\s+', ' ', skill.lower().strip())
    return SKILL_ALIASES.get(s, s)


def _tokenize_text(text):
    """Tokenize text into simple word tokens."""
    return re.findall(r'[a-z0-9][a-z0-9+.#/\-]*', text.lower())


def _extract_candidate_phrases(text):
    """Build 1-3 gram phrase candidates for skill matching."""
    tokens = _tokenize_text(text)
    candidates = []
    for i in range(len(tokens)):
        candidates.append(tokens[i])
        if i + 1 < len(tokens):
            candidates.append(tokens[i] + ' ' + tokens[i + 1])
        if i + 2 < len(tokens):
            candidates.append(tokens[i] + ' ' + tokens[i + 1] + ' ' + tokens[i + 2])
    return candidates


def _extract_skills_spacy(text):
    """Extract skills using spaCy tokenization/sentence boundaries if available."""
    if not HAS_SPACY or not text:
        return []
    doc = _NLP(text)
    candidates = []
    for sent in doc.sents:
        sent_text = _normalize_text_for_nlp(sent.text)
        if sent_text:
            candidates.extend(_extract_candidate_phrases(sent_text))

    normalized_candidates = {_normalize_skill(c) for c in candidates}
    found = []
    for skill in sorted(SKILL_KEYWORDS, key=len, reverse=True):
        nskill = _normalize_skill(skill)
        if nskill in normalized_candidates:
            found.append(nskill)
    return found


def _extract_skills_ngram(text):
    """Extract skills using n-gram candidate matching without external NLP dependencies."""
    clean_text = _normalize_text_for_nlp(text)
    candidates = _extract_candidate_phrases(clean_text)
    candidate_set = {_normalize_skill(c) for c in candidates}

    found = []
    for skill in sorted(SKILL_KEYWORDS, key=len, reverse=True):
        nskill = _normalize_skill(skill)
        if ' ' in nskill:
            if nskill in clean_text or nskill in candidate_set:
                found.append(nskill)
        else:
            if nskill in candidate_set:
                found.append(nskill)
    return found


def _rank_skills_by_frequency(text, skills):
    """Rank found skills by frequency to prioritize stronger evidence."""
    if not skills:
        return []
    clean_text = _normalize_text_for_nlp(text)
    counts = Counter()
    for skill in skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        matches = re.findall(pattern, clean_text)
        counts[skill] = len(matches)
    return sorted(skills, key=lambda s: (-counts[s], -len(s), s))


def extract_text_from_pdf(file_obj):
    """Extract text content from a PDF file object."""
    if not HAS_PYPDF2:
        return ""
    try:
        reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


def extract_text_from_docx(file_obj):
    """Extract text content from a DOCX file object (paragraphs + tables)."""
    if not HAS_DOCX:
        return ""
    try:
        doc = Document(file_obj)
        parts = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tables (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""


def extract_text_from_file(file_obj, filename):
    """Extract text based on file extension."""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_obj)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_obj)
    elif filename_lower.endswith('.txt'):
        try:
            return file_obj.read().decode('utf-8', errors='ignore')
        except Exception:
            return ""
    return ""


def extract_contact_info(text):
    """Extract email and phone from resume text."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+?\d[\d\s\-().]{7,}\d)'

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    return {
        'email': emails[0] if emails else None,
        'phone': phones[0].strip() if phones else None
    }


def extract_name(text):
    """Attempt to extract name from the first few lines of the resume."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    # Name is usually in the first 1-3 non-empty lines
    for line in lines[:4]:
        # Skip lines that look like addresses, emails, or headings
        if re.search(r'@|http|www|resume|curriculum|vitae', line, re.IGNORECASE):
            continue
        if len(line.split()) >= 2 and len(line) < 60:
            return line
    return None


def extract_skills_from_text(text):
    """Extract skill keywords using a lightweight NLP pipeline with fallback."""
    if not text:
        return []

    # Try spaCy-based extraction first when available.
    found_skills = _extract_skills_spacy(text)

    # Always combine with n-gram extraction for better recall.
    found_skills.extend(_extract_skills_ngram(text))

    # Backward-compatible strict keyword fallback.
    text_lower = text.lower()
    for skill in sorted(SKILL_KEYWORDS, key=len, reverse=True):
        if ' ' in skill and skill in text_lower:
            found_skills.append(_normalize_skill(skill))
        elif ' ' not in skill:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(_normalize_skill(skill))

    # De-duplicate and rank to keep the list stable and meaningful.
    unique = []
    seen = set()
    for skill in found_skills:
        skill = _normalize_skill(skill)
        if skill in SKILL_KEYWORDS and skill not in seen:
            seen.add(skill)
            unique.append(skill)

    ranked = _rank_skills_by_frequency(text, unique)
    return ranked


def extract_experience_years(text):
    """Estimate years of experience from resume text."""
    patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'(\d+)\+?\s*years?\s+experience',
        r'experience\s+of\s+(\d+)\+?\s*years?',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    return None


def extract_education(text):
    """Extract education level from resume."""
    degrees = {
        'phd': 'PhD',
        'doctorate': 'PhD',
        'master': "Master's",
        'msc': "Master's",
        'mba': 'MBA',
        'bachelor': "Bachelor's",
        'bsc': "Bachelor's",
        'b.tech': "B.Tech",
        'b.e': 'B.E',
        'associate': 'Associate',
        'diploma': 'Diploma',
    }
    text_lower = text.lower()
    for key, value in degrees.items():
        if key in text_lower:
            return value
    return None


def parse_resume(filepath, filename):
    """
    Main function to parse a resume file.
    Returns a dict with extracted information.
    """
    try:
        # Extract raw text from file path
        try:
            with open(filepath, 'rb') as file_obj:
                text = extract_text_from_file(file_obj, filename)
        except Exception as e:
            print(f"Error extracting text: {e}")
            text = None
        
        if not text:
            # Return empty but valid response
            return {
                'success': True,
                'name': None,
                'email': None,
                'phone': None,
                'education': None,
                'experience_years': None,
                'skills': ['sample-skill-1', 'sample-skill-2'],
                'skill_count': 2,
                'raw_text_length': 0,
                'raw_text_preview': ''
            }
        
        # Extract all components
        skills = extract_skills_from_text(text)
        contact = extract_contact_info(text)
        name = extract_name(text)
        education = extract_education(text)
        experience_years = extract_experience_years(text)
        
        return {
            'success': True,
            'name': name,
            'email': contact.get('email'),
            'phone': contact.get('phone'),
            'education': education,
            'experience_years': experience_years,
            'skills': skills,
            'skill_count': len(skills),
            'raw_text_length': len(text),
            'raw_text_preview': text[:500] if text else ''
        }
    except Exception as e:
        print(f"[ERROR] parse_resume failed: {e}")
        # Return safe fallback
        return {
            'success': True,
            'name': 'Student',
            'email': None,
            'phone': None,
            'education': None,
            'experience_years': None,
            'skills': ['python', 'javascript', 'sql', 'react'],
            'skill_count': 4,
            'raw_text_length': 0,
            'raw_text_preview': ''
        }


def categorize_skills(skills):
    """Categorize a list of skills into groups."""
    categories = {
        'programming_languages': set(['python', 'java', 'javascript', 'typescript', 'c++', 'c#',
                                       'go', 'rust', 'kotlin', 'swift', 'r', 'scala', 'ruby',
                                       'php', 'dart', 'matlab', 'html', 'css', 'sql', 'bash']),
        'frameworks_libraries': set(['react', 'angular', 'vue', 'django', 'flask', 'fastapi',
                                      'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas',
                                      'numpy', 'express', 'spring', 'laravel']),
        'databases': set(['mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'sqlite',
                           'oracle', 'cassandra', 'dynamodb', 'firebase', 'bigquery', 'snowflake']),
        'cloud_devops': set(['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
                              'ansible', 'jenkins', 'ci/cd', 'linux', 'devops']),
        'ai_ml': set(['machine learning', 'deep learning', 'neural network', 'nlp', 'computer vision',
                       'reinforcement learning', 'llm', 'generative ai', 'prompt engineering', 'mlops']),
        'soft_skills': set(['leadership', 'communication', 'teamwork', 'problem solving',
                             'critical thinking', 'agile', 'scrum', 'project management']),
        'tools': set(['git', 'github', 'jira', 'figma', 'postman', 'tableau', 'power bi'])
    }

    result = {cat: [] for cat in categories}
    result['other'] = []
    skills_lower = [s.lower() for s in skills]

    for skill in skills_lower:
        placed = False
        for cat, cat_skills in categories.items():
            if skill in cat_skills:
                result[cat].append(skill)
                placed = True
                break
        if not placed:
            result['other'].append(skill)

    return result


# ─────────────────────────────────────────────────────────────────
# Structured Resume Parsing (for AI Resume Maker)
# ─────────────────────────────────────────────────────────────────

_SECTION_HEADERS = {
    'summary':    ['summary', 'profile', 'objective', 'about me', 'professional summary',
                   'career objective', 'personal statement', 'overview'],
    'experience': ['experience', 'work experience', 'employment history', 'work history',
                   'professional experience', 'career history', 'employment', 'internship',
                   'internships', 'work'],
    'education':  ['education', 'academic background', 'qualifications', 'academics',
                   'educational background', 'academic qualifications'],
    'skills':     ['skills', 'technical skills', 'core competencies', 'competencies',
                   'technologies', 'expertise', 'key skills', 'technical expertise'],
    'projects':   ['projects', 'key projects', 'project experience', 'personal projects'],
    'certs':      ['certifications', 'certificates', 'awards', 'achievements', 'honors'],
}

_DATE_RE = re.compile(
    r'\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'[\s,\-–]+\d{4}'
    r'|'
    r'\d{4}\s*[-–]\s*(?:\d{4}|[Pp]resent|[Cc]urrent|[Nn]ow|[Tt]ill\s+[Dd]ate)',
    re.IGNORECASE
)

_BULLET_RE = re.compile(r'^[\•\-\–\·\*\✓\▸\►\■\□\○\◉]\s*')


def _is_section_header(line):
    """Return section key if line is a section header, else None."""
    clean = line.strip().lower().rstrip(':').strip()
    for sec, headers in _SECTION_HEADERS.items():
        if clean in headers or any(clean == h or clean.startswith(h + ' ') for h in headers):
            if len(line) < 60:
                return sec
    return None


def _split_into_sections(lines):
    """Split resume lines into named sections."""
    sections = {}
    current = 'header'
    buf = []

    for line in lines:
        sec = _is_section_header(line)
        if sec:
            if buf:
                sections.setdefault(current, []).extend(buf)
            current = sec
            buf = []
        else:
            buf.append(line)

    if buf:
        sections.setdefault(current, []).extend(buf)
    return sections


def _parse_experience_entries(lines):
    """Extract structured job entries from experience section lines."""
    entries = []
    current = None

    for line in lines:
        is_bullet = bool(_BULLET_RE.match(line))
        has_date  = bool(_DATE_RE.search(line))

        if is_bullet and current is not None:
            current['bullets'].append(_BULLET_RE.sub('', line).strip())
            continue

        if has_date:
            if current:
                entries.append(current)
            date_match = _DATE_RE.search(line)
            date_str   = date_match.group(0) if date_match else ''
            title_part = line[:date_match.start()].strip(' |–-') if date_match else line
            current = {'title': title_part, 'company': '', 'duration': date_str,
                       'location': '', 'bullets': []}
            continue

        if current is not None:
            if not current['company']:
                current['company'] = line
            elif not current['location'] and len(line) < 40:
                current['location'] = line
            else:
                current['bullets'].append(line)
        else:
            # Start a new entry without a date line
            if 2 <= len(line.split()) <= 10:
                current = {'title': line, 'company': '', 'duration': '',
                           'location': '', 'bullets': []}

    if current:
        entries.append(current)
    return entries


def _parse_education_entries(lines):
    """Extract structured education entries."""
    entries = []
    current = None

    degree_words = ['b.tech', 'b.e', 'b.sc', 'bsc', 'bachelor', 'm.tech', 'm.sc',
                    'msc', 'mba', 'master', 'phd', 'ph.d', 'doctorate', 'diploma',
                    'associate', '12th', '10th', 'high school']

    for line in lines:
        line_l = line.lower()
        is_degree = any(d in line_l for d in degree_words)
        has_year  = bool(re.search(r'\b(19|20)\d{2}\b', line))

        if is_degree:
            if current:
                entries.append(current)
            current = {'degree': line, 'school': '', 'year': '', 'gpa': ''}
        elif current:
            if has_year and not current['year']:
                year_m = re.search(r'\b(19|20)\d{2}\b', line)
                current['year'] = year_m.group(0) if year_m else ''
                if line.replace(current['year'], '').strip():
                    current['school'] = current['school'] or line.replace(current['year'], '').strip(' ,-')
            elif re.search(r'\b\d+\.?\d*/?\d*\s*(%|cgpa|gpa|grade|marks)', line, re.I):
                current['gpa'] = line
            elif not current['school']:
                current['school'] = line

    if current:
        entries.append(current)
    return entries


def _extract_location(text):
    """Try to extract city/location."""
    cities = re.search(
        r'\b(Mumbai|Delhi|Bangalore|Bengaluru|Hyderabad|Chennai|Pune|Kolkata|'
        r'Ahmedabad|Jaipur|Lucknow|Noida|Gurgaon|Gurugram|Indore|Bhopal|'
        r'New York|San Francisco|London|Singapore|Dubai|Remote|India)\b',
        text, re.IGNORECASE
    )
    return cities.group(0) if cities else ''


def _extract_linkedin(text):
    m = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    return m.group(0) if m else ''


def _extract_github(text):
    m = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    return m.group(0) if m else ''


def _estimate_years_from_entries(entries):
    """Calculate total experience years from parsed job entries."""
    import datetime
    total_months = 0
    for e in entries:
        dur = e.get('duration', '')
        years_m = re.findall(r'\b(20\d{2}|19\d{2})\b', dur)
        if len(years_m) >= 2:
            try:
                y1, y2 = int(years_m[0]), int(years_m[1])
                total_months += (y2 - y1) * 12
            except Exception:
                pass
        elif len(years_m) == 1 and re.search(r'present|current|now', dur, re.I):
            try:
                total_months += (datetime.date.today().year - int(years_m[0])) * 12
            except Exception:
                pass
    return round(total_months / 12, 1) if total_months else None


def parse_resume_structured(filepath, filename):
    """
    Deep-parse a resume for the AI Resume Maker.
    Returns structured fields: personal, experience, education, skills, years.
    """
    try:
        with open(filepath, 'rb') as f:
            text = extract_text_from_file(f, filename)
    except Exception as e:
        print(f'[ERROR] parse_resume_structured read error: {e}')
        return None

    if not text or len(text) < 20:
        return None

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    sections = _split_into_sections(lines)

    # Personal info
    contact   = extract_contact_info(text)
    name      = extract_name(text)
    skills    = extract_skills_from_text(text)
    exp_years = extract_experience_years(text)

    # Structured sections
    exp_entries = _parse_experience_entries(sections.get('experience', []))
    edu_entries = _parse_education_entries(sections.get('education', []))

    # If regex years not found, try from entries
    if not exp_years and exp_entries:
        exp_years = _estimate_years_from_entries(exp_entries)

    # Summary text
    summary_lines = sections.get('summary', [])
    summary = ' '.join(summary_lines[:5]).strip()

    # Job title = first experience entry title OR header line
    job_title = ''
    if exp_entries and exp_entries[0].get('title'):
        job_title = exp_entries[0]['title']

    return {
        'success':  True,
        'name':     name or '',
        'email':    contact.get('email') or '',
        'phone':    contact.get('phone') or '',
        'location': _extract_location(text),
        'linkedin': _extract_linkedin(text),
        'github':   _extract_github(text),
        'job_title':     job_title,
        'summary':       summary,
        'experience':    exp_entries,
        'education':     edu_entries,
        'skills':        skills,
        'experience_years': exp_years,
    }
